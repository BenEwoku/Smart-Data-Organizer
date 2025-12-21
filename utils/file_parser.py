"""
File parsing utilities for PDF, DOCX, and Excel files
Extracts tables and text data for conversion
"""

import pandas as pd
import streamlit as st
from io import BytesIO
import re

def parse_uploaded_file(uploaded_file):
    """
    Parse uploaded file based on type
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        pd.DataFrame (ALWAYS returns DataFrame, never None)
    """
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    # Initialize default result
    default_df = pd.DataFrame({"File Info": [
        f"File: {uploaded_file.name}",
        f"Type: {file_ext.upper()}",
        "Status: Processing..."
    ]})
    
    try:
        if file_ext == 'csv':
            df = parse_csv(uploaded_file)
        elif file_ext == 'txt':
            df = parse_txt(uploaded_file)
        elif file_ext in ['xlsx', 'xls']:
            df = parse_excel(uploaded_file)
        elif file_ext == 'pdf':
            df = parse_pdf(uploaded_file)
        elif file_ext in ['docx', 'doc']:
            df = parse_docx(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_ext}")
            return pd.DataFrame({"Error": [f"Unsupported file type: {file_ext}"]})
        
        # FINAL SAFETY CHECK: Ensure we never return None
        if df is None:
            st.warning(f"Parser returned None for {file_ext} file")
            return default_df
            
        # Ensure it's a DataFrame
        if not isinstance(df, pd.DataFrame):
            st.warning(f"Parser returned {type(df)} instead of DataFrame")
            return default_df
            
        # Ensure it's not empty
        if df.empty:
            st.info(f"File parsed but returned empty DataFrame")
            # Add at least one row
            if len(df) == 0:
                df = pd.DataFrame({"Info": ["File parsed successfully but no data extracted"]})
        
        return df
        
    except Exception as e:
        st.error(f"Unexpected error parsing {file_ext.upper()} file: {str(e)}")
        # Always return a DataFrame
        return pd.DataFrame({
            "Error": [f"Unexpected error: {str(e)[:200]}"],
            "File": [uploaded_file.name],
            "Type": [file_ext]
        })

def parse_csv(file):
    """Parse CSV file - always returns DataFrame"""
    try:
        df = pd.read_csv(file)
        return df
    except Exception as e:
        # Try different encodings
        try:
            df = pd.read_csv(file, encoding='latin-1')
            return df
        except:
            st.error(f"CSV parsing error: {str(e)}")
            # Return empty DataFrame instead of None
            return pd.DataFrame({"Error": [f"CSV parsing failed: {str(e)[:100]}"]})

def parse_excel(file):
    """Parse Excel file - always returns DataFrame"""
    try:
        # Try openpyxl first (for .xlsx)
        try:
            df = pd.read_excel(file, engine='openpyxl')
            return df
        except:
            # Fall back to xlrd (for .xls)
            df = pd.read_excel(file, engine='xlrd')
            return df
    except Exception as e:
        st.error(f"Excel parsing error: {str(e)}")
        # Return empty DataFrame instead of None
        return pd.DataFrame({"Error": [f"Excel parsing failed: {str(e)[:100]}"]})

def parse_txt(file):
    """Parse TXT file - always returns DataFrame"""
    try:
        content = file.read().decode('utf-8')
        
        # Try to detect structure
        from utils.parser import parse_text_to_dataframe
        df = parse_text_to_dataframe(content)
        
        if df is not None:
            return df
        else:
            return pd.DataFrame({"Content": [content[:500] + "..." if len(content) > 500 else content]})
    except Exception as e:
        st.error(f"TXT parsing error: {str(e)}")
        # Return empty DataFrame instead of None
        return pd.DataFrame({"Error": [f"TXT parsing failed: {str(e)[:100]}"]})

def parse_pdf(file):
    """
    Parse PDF file and extract tables
    
    Returns:
        DataFrame (always returns a DataFrame, even if empty)
    """
    st.info("Extracting data from PDF...")
    
    tables = []
    
    # Method 1: Try pdfplumber (best for most PDFs)
    try:
        import pdfplumber
        
        with pdfplumber.open(file) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_tables = page.extract_tables()
                
                if page_tables:
                    for table in page_tables:
                        if table and len(table) > 1:
                            # Convert to DataFrame
                            # Handle empty headers
                            headers = table[0] if table[0] else [f"Column_{i}" for i in range(len(table[0]))]
                            df = pd.DataFrame(table[1:], columns=headers)
                            df = df.dropna(how='all', axis=1)  # Remove empty columns
                            df = df.dropna(how='all', axis=0)  # Remove empty rows
                            
                            if len(df) > 0:
                                # Clean column names immediately
                                df.columns = [f'Column_{i}' if pd.isna(col) else str(col) for i, col in enumerate(df.columns)]
                                tables.append(df)
                                st.success(f"Found table on page {page_num}")
        
        if tables:
            if len(tables) == 1:
                return tables[0]
            else:
                st.info(f"Found {len(tables)} tables. Combining them...")
                return combine_tables(tables)
    
    except ImportError:
        st.warning("pdfplumber not installed. Trying alternative method...")
    except Exception as e:
        st.warning(f"pdfplumber extraction failed: {str(e)}")
    
    # Method 2: Try tabula-py (good for complex PDFs)
    try:
        import tabula
        
        # Reset file pointer
        file.seek(0)
        
        # Extract tables
        tables = tabula.read_pdf(file, pages='all', multiple_tables=True)
        
        if tables:
            st.success(f"Found {len(tables)} table(s) using Tabula")
            
            # Clean column names for all tables
            for i, df in enumerate(tables):
                df.columns = [f'Column_{j}' if pd.isna(col) else str(col) for j, col in enumerate(df.columns)]
            
            if len(tables) == 1:
                return tables[0]
            else:
                return combine_tables(tables)
    
    except ImportError:
        st.warning("tabula-py not installed. Trying text extraction...")
    except Exception as e:
        st.warning(f"Tabula extraction failed: {str(e)}")
    
    # Method 3: Extract text and try to parse
    try:
        import PyPDF2
        
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        
        all_text = ""
        for page in pdf_reader.pages:
            all_text += page.extract_text() + "\n"
        
        if all_text.strip():
            st.info("No tables found. Attempting to parse text content...")
            
            from utils.parser import parse_text_to_dataframe
            df = parse_text_to_dataframe(all_text)
            
            if df is not None and len(df) > 0:
                st.success("Successfully parsed text content")
                # Clean column names
                df.columns = [f'Column_{i}' if pd.isna(col) else str(col) for i, col in enumerate(df.columns)]
                return df
    
    except Exception as e:
        st.warning(f"Text extraction failed: {str(e)}")
    
    # CRITICAL: ALWAYS return a DataFrame, never None
    st.warning("Could not extract structured data from PDF. Creating empty DataFrame.")
    
    # Create empty DataFrame with fallback column
    empty_df = pd.DataFrame({"Empty": ["No data extracted from PDF"]})
    
    # Show troubleshooting tips
    with st.expander("Troubleshooting PDF extraction"):
        st.markdown("""
        **Tips for better PDF parsing:**
        - Ensure PDF contains actual tables (not images)
        - Use PDFs with clear table borders
        - Consider using OCR for scanned documents
        - Try converting PDF to CSV/Excel first for better results
        """)
    
    return empty_df

def parse_docx(file):
    """
    Parse DOCX file and extract tables
    
    Returns:
        DataFrame (never returns None)
    """
    try:
        from docx import Document
        
        doc = Document(file)
        tables = []
        
        st.info("Extracting tables from Word document...")
        
        # Extract all tables
        for table_num, table in enumerate(doc.tables, 1):
            data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                data.append(row_data)
            
            if len(data) > 1:
                # First row as header
                # Handle empty headers
                headers = data[0] if data[0] else [f"Column_{i}" for i in range(len(data[0]))]
                df = pd.DataFrame(data[1:], columns=headers)
                df = df.dropna(how='all', axis=1)
                df = df.dropna(how='all', axis=0)
                
                if len(df) > 0:
                    # Clean column names immediately
                    df.columns = [f'Column_{i}' if pd.isna(col) else str(col) for i, col in enumerate(df.columns)]
                    tables.append(df)
                    st.success(f"Found table {table_num}")
        
        if tables:
            if len(tables) == 1:
                return tables[0]
            else:
                st.info(f"Found {len(tables)} tables. Combining them...")
                return combine_tables(tables)
        else:
            # No tables found, try to extract text
            st.info("No tables found. Extracting text content...")
            
            all_text = ""
            for para in doc.paragraphs:
                all_text += para.text + "\n"
            
            if all_text.strip():
                from utils.parser import parse_text_to_dataframe
                df = parse_text_to_dataframe(all_text)
                
                if df is not None and len(df) > 0:
                    st.success("Successfully parsed text content")
                    # Clean column names
                    df.columns = [f'Column_{i}' if pd.isna(col) else str(col) for i, col in enumerate(df.columns)]
                    return df
            
            # No data found - return empty DataFrame instead of None
            st.warning("No structured data found in document. Creating empty DataFrame.")
            return pd.DataFrame({"Info": ["No structured data found in Word document"]})
    
    except ImportError:
        st.error("python-docx not installed. Cannot parse DOCX files.")
        # Return empty DataFrame instead of None
        return pd.DataFrame({"Error": ["python-docx library not installed"]})
    except Exception as e:
        st.error(f"DOCX parsing error: {str(e)}")
        # Return empty DataFrame instead of None
        return pd.DataFrame({"Error": [f"DOCX parsing failed: {str(e)[:100]}"]})

def combine_tables(tables):
    """
    Combine multiple DataFrames intelligently
    
    Args:
        tables: List of DataFrames
        
    Returns:
        pd.DataFrame: Combined DataFrame
    """
    if not tables:
        return None
    
    if len(tables) == 1:
        return tables[0]
    
    st.write("### Multiple Tables Found")
    
    # Show preview of each table
    for i, df in enumerate(tables, 1):
        with st.expander(f"Table {i} - {len(df)} rows × {len(df.columns)} columns"):
            st.dataframe(df.head())
    
    # Ask user how to combine
    combine_method = st.radio(
        "How should we combine these tables?",
        [
            "Use first table only",
            "Concatenate all tables (stack vertically)",
            "Merge all tables (combine columns)",
            "Keep separate - export individually"
        ]
    )
    
    if combine_method == "Use first table only":
        return tables[0]
    
    elif combine_method == "Concatenate all tables (stack vertically)":
        try:
            # Try to align columns
            combined = pd.concat(tables, ignore_index=True)
            st.success(f"✅ Combined {len(tables)} tables")
            return combined
        except:
            st.error("Could not concatenate tables (different structures)")
            return tables[0]
    
    elif combine_method == "Merge all tables (combine columns)":
        try:
            combined = tables[0]
            for df in tables[1:]:
                combined = pd.concat([combined, df], axis=1)
            st.success(f"✅ Merged {len(tables)} tables")
            return combined
        except:
            st.error("Could not merge tables")
            return tables[0]
    
    else:  # Keep separate
        st.info("Tables will be processed separately")
        # For now, return first table
        # In future, could allow user to select which one
        return tables[0]

def extract_text_from_pdf(file):
    """Extract raw text from PDF (no table detection)"""
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text()
        
        return text
    except:
        return None

def extract_text_from_docx(file):
    """Extract raw text from DOCX"""
    try:
        from docx import Document
        
        doc = Document(file)
        text = ""
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        return text
    except:
        return None